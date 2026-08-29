from __future__ import annotations

from dataclasses import dataclass, field
from io import BytesIO
from pathlib import Path
import re
from typing import Any, Iterable


class DocumentExtractionError(RuntimeError):
    pass


@dataclass
class ExtractedPage:
    source_name: str
    page: int | None
    text: str
    method: str

    def tagged_text(self) -> str:
        page = "-" if self.page is None else str(self.page)
        return f"[SOURCE {self.source_name} | PAGE {page} | METHOD {self.method}]\n{self.text.strip()}"


@dataclass
class DocumentExtractionResult:
    pages: list[ExtractedPage] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)

    @property
    def text(self) -> str:
        return "\n\n".join(p.tagged_text() for p in self.pages if p.text.strip())

    @property
    def source_names(self) -> list[str]:
        seen: list[str] = []
        for p in self.pages:
            if p.source_name not in seen:
                seen.append(p.source_name)
        return seen


def _read_bytes(item: Any) -> tuple[str, bytes]:
    if isinstance(item, (str, Path)):
        p = Path(item)
        return p.name, p.read_bytes()
    if isinstance(item, tuple) and len(item) == 2:
        name, data = item
        return str(name), bytes(data)
    name = getattr(item, "name", None)
    if name and hasattr(item, "getvalue"):
        return str(name), bytes(item.getvalue())
    if name and hasattr(item, "read"):
        pos = None
        try:
            pos = item.tell()
        except Exception:
            pass
        data = item.read()
        if pos is not None:
            try:
                item.seek(pos)
            except Exception:
                pass
        return str(name), bytes(data)
    raise DocumentExtractionError(
        "Dokumen harus berupa path, (filename, bytes), atau file-like object yang memiliki name/read()."
    )


def _extract_pdf(name: str, data: bytes, out: DocumentExtractionResult) -> None:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise DocumentExtractionError("Install pypdf untuk ekstraksi PDF.") from exc

    reader = PdfReader(BytesIO(data))
    for i, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        if len(text) < 20:
            out.warnings.append(
                f"{name} halaman {i}: teks sangat sedikit/kosong. Kemungkinan scanned PDF; "
                "jalankan OCR/VLM sebelum scoring dan jangan anggap halaman ini sudah terbaca."
            )
        out.pages.append(ExtractedPage(name, i, text, "pypdf"))


def _extract_docx(name: str, data: bytes, out: DocumentExtractionResult) -> None:
    try:
        from docx import Document
    except ImportError as exc:
        raise DocumentExtractionError("Install python-docx untuk ekstraksi DOCX.") from exc
    doc = Document(BytesIO(data))
    blocks: list[str] = []
    for p in doc.paragraphs:
        if p.text.strip():
            blocks.append(p.text.strip())
    for table in doc.tables:
        for row in table.rows:
            blocks.append(" | ".join(cell.text.strip() for cell in row.cells))
    out.pages.append(ExtractedPage(name, None, "\n".join(blocks), "python-docx"))


def _extract_textlike(name: str, data: bytes, out: DocumentExtractionResult) -> None:
    text = data.decode("utf-8", errors="replace")
    out.pages.append(ExtractedPage(name, None, text, "utf-8"))


def extract_documents(items: Iterable[Any]) -> DocumentExtractionResult:
    """Extract machine-readable text from common RM input files.

    Supported here: PDF (text layer), DOCX, TXT/MD/JSON/CSV. Image/scanned-PDF OCR
    is intentionally not guessed. The result raises a warning so an OCR/VLM adapter
    can be inserted explicitly before risk scoring.
    """
    out = DocumentExtractionResult()
    for item in items:
        name, data = _read_bytes(item)
        ext = Path(name).suffix.lower()
        if ext == ".pdf":
            _extract_pdf(name, data, out)
        elif ext == ".docx":
            _extract_docx(name, data, out)
        elif ext in {".txt", ".md", ".json", ".csv"}:
            _extract_textlike(name, data, out)
        elif ext in {".png", ".jpg", ".jpeg", ".webp", ".tif", ".tiff"}:
            out.warnings.append(
                f"{name}: file image membutuhkan OCR/VLM adapter; belum diekstrak oleh text-layer extractor."
            )
        else:
            out.warnings.append(f"{name}: format {ext or '(tanpa ekstensi)'} belum didukung extractor.")
    return out



def _clean_ocr_text(text: str) -> str:
    import re
    if not text:
        return ""
    text = text.replace("\x0c", "")
    lines = [line.rstrip() for line in text.splitlines()]
    return re.sub(r"\n{3,}", "\n\n", "\n".join(lines)).strip()


def _prepare_image_for_ocr(image_bytes: bytes):
    try:
        from PIL import Image, ImageOps
    except ImportError as exc:
        raise DocumentExtractionError("Install pillow untuk OCR image.") from exc
    image = Image.open(BytesIO(image_bytes)).convert("L")
    image = ImageOps.autocontrast(image)
    if image.width < 1500:
        scale = 1500 / max(image.width, 1)
        image = image.resize((int(image.width * scale), int(image.height * scale)))
    return image


def _run_tesseract(image_bytes: bytes, language: str = "ind+eng") -> str:
    try:
        import pytesseract
    except ImportError as exc:
        raise DocumentExtractionError("Install pytesseract untuk OCR.") from exc
    image = _prepare_image_for_ocr(image_bytes)
    try:
        text = pytesseract.image_to_string(image, lang=language, config="--oem 3 --psm 3")
    except Exception:
        text = pytesseract.image_to_string(image, lang="eng", config="--oem 3 --psm 3")
    return _clean_ocr_text(text)


def _run_vlm_transcription(
    image_bytes: bytes,
    *,
    source_name: str,
    page_number: int | None,
    ollama_url: str,
    vlm_model: str,
    timeout: int = 300,
) -> str:
    import base64
    try:
        import requests
    except ImportError as exc:
        raise DocumentExtractionError("Install requests untuk VLM fallback.") from exc

    location = f"halaman {page_number}" if page_number is not None else "image"
    prompt = f"""
Anda adalah document transcription engine untuk dokumen kredit/perbankan.
Dokumen: {source_name}
Lokasi: {location}

Transkripsikan SEMUA teks yang benar-benar terlihat pada gambar.
Jangan membuat ringkasan, analisis kredit, atau menebak angka yang tidak terbaca.
Pertahankan rupiah, persen, tanggal, nama perusahaan, nama fasilitas, angka laporan
keuangan, tanda negatif, dan desimal. Jika ada tabel, pertahankan hubungan baris/kolom.
Jika bagian tidak terbaca tulis [TIDAK TERBACA]. Keluarkan HANYA hasil transkripsi.
""".strip()

    payload = {
        "model": vlm_model,
        "messages": [{
            "role": "user",
            "content": prompt,
            "images": [base64.b64encode(image_bytes).decode("utf-8")],
        }],
        "stream": False,
        "options": {"temperature": 0.0},
        "keep_alive": "5m",
    }
    r = requests.post(f"{ollama_url.rstrip('/')}/api/chat", json=payload, timeout=timeout)
    r.raise_for_status()
    return _clean_ocr_text(r.json().get("message", {}).get("content", ""))


def extract_documents_multimodal(
    items: Iterable[Any],
    *,
    ollama_url: str = "http://127.0.0.1:11434",
    vlm_model: str = "qwen3-vl:4b-instruct",
    min_native_chars: int = 80,
    min_ocr_chars: int = 80,
    pdf_dpi: int = 200,
    ocr_language: str = "ind+eng",
    use_vlm_fallback: bool = True,
) -> DocumentExtractionResult:
    """Native text first, Tesseract OCR second, VLM transcription last.

    Multiple documents are supported. VLM is only called for pages/images whose native
    text and Tesseract OCR are both too sparse. The VLM is a transcription layer only;
    it does not calculate model features or risk values.
    """
    buffered = [_read_bytes(item) for item in items]
    base = extract_documents(buffered)
    out = DocumentExtractionResult(pages=list(base.pages), warnings=list(base.warnings))

    try:
        import fitz  # PyMuPDF
    except ImportError as exc:
        raise DocumentExtractionError("Install pymupdf untuk render scanned PDF.") from exc

    image_exts = {".png", ".jpg", ".jpeg", ".webp", ".tif", ".tiff"}

    for name, data in buffered:
        ext = Path(name).suffix.lower()
        if ext == ".pdf":
            pdf = fitz.open(stream=data, filetype="pdf")
            for idx in range(len(pdf)):
                page_no = idx + 1
                existing = next(
                    (x for x in out.pages if x.source_name == name and x.page == page_no),
                    None,
                )
                native = existing.text.strip() if existing else ""
                if len(native) >= min_native_chars:
                    continue

                zoom = pdf_dpi / 72
                pix = pdf[idx].get_pixmap(matrix=fitz.Matrix(zoom, zoom), alpha=False)
                image_bytes = pix.tobytes("png")
                try:
                    ocr = _run_tesseract(image_bytes, ocr_language)
                except Exception as exc:
                    ocr = ""
                    out.warnings.append(f"{name} halaman {page_no}: OCR gagal: {exc}")

                final, method = ocr, "tesseract-ocr"
                if len(ocr) < min_ocr_chars and use_vlm_fallback:
                    try:
                        vlm = _run_vlm_transcription(
                            image_bytes,
                            source_name=name,
                            page_number=page_no,
                            ollama_url=ollama_url,
                            vlm_model=vlm_model,
                        )
                    except Exception as exc:
                        vlm = ""
                        out.warnings.append(f"{name} halaman {page_no}: VLM gagal: {exc}")
                    if vlm:
                        final, method = vlm, f"vlm:{vlm_model}"
                    elif ocr:
                        method = "tesseract-ocr-low-text"
                    else:
                        final, method = native, "pypdf-low-text"

                if existing:
                    existing.text, existing.method = final, method
                else:
                    out.pages.append(ExtractedPage(name, page_no, final, method))
            pdf.close()

        elif ext in image_exts:
            # Remove text-layer placeholder warning for this image because we now process it.
            out.warnings = [w for w in out.warnings if not w.startswith(f"{name}: file image")]
            try:
                ocr = _run_tesseract(data, ocr_language)
            except Exception as exc:
                ocr = ""
                out.warnings.append(f"{name}: OCR gagal: {exc}")
            final, method = ocr, "tesseract-ocr"
            if len(ocr) < min_ocr_chars and use_vlm_fallback:
                try:
                    vlm = _run_vlm_transcription(
                        data,
                        source_name=name,
                        page_number=None,
                        ollama_url=ollama_url,
                        vlm_model=vlm_model,
                    )
                except Exception as exc:
                    vlm = ""
                    out.warnings.append(f"{name}: VLM gagal: {exc}")
                if vlm:
                    final, method = vlm, f"vlm:{vlm_model}"
                elif ocr:
                    method = "tesseract-ocr-low-text"
            out.pages.append(ExtractedPage(name, None, final, method))

    # Remove stale scanned-PDF warnings for pages that were successfully enriched.
    cleaned_warnings = []
    for w in out.warnings:
        if "teks sangat sedikit/kosong" in w:
            m = re.search(r"^(.*?) halaman (\d+):", w)
            if m:
                src, page = m.group(1), int(m.group(2))
                page_obj = next((x for x in out.pages if x.source_name == src and x.page == page), None)
                if page_obj and len(page_obj.text.strip()) >= min_ocr_chars:
                    continue
        cleaned_warnings.append(w)
    out.warnings = cleaned_warnings
    out.pages.sort(key=lambda x: (x.source_name, x.page or 0))
    return out
